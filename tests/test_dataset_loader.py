import json

import pytest
from docx import Document

from src import dataset_loader
from src.dataset_loader import (
    DatasetSchemaError,
    UnsupportedDatasetFormatError,
    load_dataset,
)


DATASET_SCHEMA = {
    "id": int,
    "prompt": str,
    "expected": str,
}


def _write_docx_table(path, headers, rows):
    document = Document()
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value

    document.save(path)


@pytest.mark.parametrize("extension", [".csv", ".json", ".docx"])
def test_load_dataset_supports_csv_json_and_docx_with_schema_validation(
    tmp_path,
    extension,
):
    dataset_path = tmp_path / f"benchmark{extension}"

    if extension == ".csv":
        dataset_path.write_text(
            "id,prompt,expected\n"
            "1,Summarize invoice INV-001,Invoice INV-001 total is 1500000\n"
            "2,Extract customer name,Acme Corp\n",
            encoding="utf-8",
        )
    elif extension == ".json":
        dataset_path.write_text(
            json.dumps(
                [
                    {
                        "id": 1,
                        "prompt": "Summarize invoice INV-001",
                        "expected": "Invoice INV-001 total is 1500000",
                    },
                    {
                        "id": 2,
                        "prompt": "Extract customer name",
                        "expected": "Acme Corp",
                    },
                ]
            ),
            encoding="utf-8",
        )
    else:
        _write_docx_table(
            dataset_path,
            headers=["id", "prompt", "expected"],
            rows=[
                [
                    "1",
                    "Summarize invoice INV-001",
                    "Invoice INV-001 total is 1500000",
                ],
                ["2", "Extract customer name", "Acme Corp"],
            ],
        )

    rows = load_dataset(dataset_path, schema=DATASET_SCHEMA)

    assert rows == [
        {
            "id": 1,
            "prompt": "Summarize invoice INV-001",
            "expected": "Invoice INV-001 total is 1500000",
        },
        {
            "id": 2,
            "prompt": "Extract customer name",
            "expected": "Acme Corp",
        },
    ]
    assert all(isinstance(row["id"], int) for row in rows)


def test_load_dataset_rejects_unrecognized_file_format(tmp_path):
    dataset_path = tmp_path / "benchmark.txt"
    dataset_path.write_text("id,prompt,expected\n1,Prompt,Expected\n", encoding="utf-8")

    with pytest.raises(
        UnsupportedDatasetFormatError,
        match=r"Unsupported dataset format.*\.txt",
    ):
        load_dataset(dataset_path, schema=DATASET_SCHEMA)


def test_load_dataset_rejects_empty_schema(tmp_path):
    dataset_path = tmp_path / "benchmark.csv"
    dataset_path.write_text("id,prompt,expected\n1,Prompt,Expected\n", encoding="utf-8")

    with pytest.raises(DatasetSchemaError, match="Schema must define"):
        load_dataset(dataset_path, schema={})


def test_load_dataset_rejects_schema_with_non_type_value(tmp_path):
    dataset_path = tmp_path / "benchmark.csv"
    dataset_path.write_text("id,prompt,expected\n1,Prompt,Expected\n", encoding="utf-8")

    with pytest.raises(DatasetSchemaError, match="prompt"):
        load_dataset(dataset_path, schema={"id": int, "prompt": "str"})


def test_load_dataset_rejects_missing_required_columns(tmp_path):
    dataset_path = tmp_path / "benchmark.csv"
    dataset_path.write_text(
        "id,prompt\n"
        "1,Summarize invoice INV-001\n",
        encoding="utf-8",
    )

    with pytest.raises(DatasetSchemaError, match="Missing required columns.*expected"):
        load_dataset(dataset_path, schema=DATASET_SCHEMA)


def test_load_dataset_rejects_values_that_do_not_match_schema_type(tmp_path):
    dataset_path = tmp_path / "benchmark.csv"
    dataset_path.write_text(
        "id,prompt,expected\n"
        "not-an-integer,Summarize invoice INV-001,Invoice INV-001 total is 1500000\n",
        encoding="utf-8",
    )

    with pytest.raises(DatasetSchemaError, match="id.*int"):
        load_dataset(dataset_path, schema=DATASET_SCHEMA)


def test_load_dataset_rejects_empty_required_values(tmp_path):
    dataset_path = tmp_path / "benchmark.csv"
    dataset_path.write_text(
        "id,prompt,expected\n"
        "1,,Expected answer\n",
        encoding="utf-8",
    )

    with pytest.raises(DatasetSchemaError, match="prompt.*required"):
        load_dataset(dataset_path, schema=DATASET_SCHEMA)


def test_load_dataset_casts_non_string_values_to_string_for_string_columns(tmp_path):
    dataset_path = tmp_path / "benchmark.json"
    dataset_path.write_text(
        json.dumps(
            [
                {
                    "id": 1,
                    "prompt": 12345,
                    "expected": 67890,
                }
            ]
        ),
        encoding="utf-8",
    )

    rows = load_dataset(dataset_path, schema=DATASET_SCHEMA)

    assert rows == [{"id": 1, "prompt": "12345", "expected": "67890"}]


def test_load_dataset_supports_wrapped_json_rows_and_normalizes_keys(tmp_path):
    dataset_path = tmp_path / "benchmark.json"
    dataset_path.write_text(
        json.dumps(
            {
                "rows": [
                    {
                        "id ": 1,
                        " prompt ": "Summarize invoice INV-001",
                        "expected ": "Invoice INV-001 total is 1500000",
                    }
                ]
            }
        ),
        encoding="utf-8",
    )

    rows = load_dataset(dataset_path, schema=DATASET_SCHEMA)

    assert rows == [
        {
            "id": 1,
            "prompt": "Summarize invoice INV-001",
            "expected": "Invoice INV-001 total is 1500000",
        }
    ]


def test_load_dataset_returns_empty_rows_for_empty_json_wrapper(tmp_path):
    dataset_path = tmp_path / "benchmark.json"
    dataset_path.write_text(json.dumps({"rows": []}), encoding="utf-8")

    rows = load_dataset(dataset_path, schema=DATASET_SCHEMA)

    assert rows == []


def test_load_dataset_supports_wrapped_json_data_key(tmp_path):
    dataset_path = tmp_path / "benchmark.json"
    dataset_path.write_text(
        json.dumps(
            {
                "data": [
                    {
                        "id": 1,
                        "prompt": "Prompt from data wrapper",
                        "expected": "Expected from data wrapper",
                    }
                ]
            }
        ),
        encoding="utf-8",
    )

    rows = load_dataset(dataset_path, schema=DATASET_SCHEMA)

    assert rows == [
        {
            "id": 1,
            "prompt": "Prompt from data wrapper",
            "expected": "Expected from data wrapper",
        }
    ]


def test_load_dataset_supports_single_json_object(tmp_path):
    dataset_path = tmp_path / "benchmark.json"
    dataset_path.write_text(
        json.dumps(
            {
                "id": 1,
                "prompt": "Single object prompt",
                "expected": "Single object expected",
            }
        ),
        encoding="utf-8",
    )

    rows = load_dataset(dataset_path, schema=DATASET_SCHEMA)

    assert rows == [
        {
            "id": 1,
            "prompt": "Single object prompt",
            "expected": "Single object expected",
        }
    ]


def test_load_dataset_rejects_json_wrapper_that_is_not_a_list(tmp_path):
    dataset_path = tmp_path / "benchmark.json"
    dataset_path.write_text(
        json.dumps({"rows": {"id": 1, "prompt": "Prompt", "expected": "Expected"}}),
        encoding="utf-8",
    )

    with pytest.raises(DatasetSchemaError, match="object or list of objects"):
        load_dataset(dataset_path, schema=DATASET_SCHEMA)


def test_load_dataset_rejects_json_non_object_before_first_record(tmp_path):
    dataset_path = tmp_path / "benchmark.json"
    dataset_path.write_text(json.dumps(["not-a-row"]), encoding="utf-8")

    with pytest.raises(DatasetSchemaError, match="rows must be objects"):
        load_dataset(dataset_path, schema=DATASET_SCHEMA)


def test_load_dataset_rejects_json_non_object_after_first_record(tmp_path):
    dataset_path = tmp_path / "benchmark.json"
    dataset_path.write_text(
        json.dumps(
            [
                {"id": 1, "prompt": "Prompt", "expected": "Expected"},
                "not-a-row",
            ]
        ),
        encoding="utf-8",
    )

    with pytest.raises(DatasetSchemaError, match="rows must be objects"):
        load_dataset(dataset_path, schema=DATASET_SCHEMA)


def test_load_dataset_skips_empty_json_records_around_valid_records(tmp_path):
    dataset_path = tmp_path / "benchmark.json"
    dataset_path.write_text(
        json.dumps(
            [
                {},
                {"id": 1, "prompt": "Prompt", "expected": "Expected"},
                {},
            ]
        ),
        encoding="utf-8",
    )

    rows = load_dataset(dataset_path, schema=DATASET_SCHEMA)

    assert rows == [{"id": 1, "prompt": "Prompt", "expected": "Expected"}]


def test_load_dataset_rejects_json_with_only_empty_records(tmp_path):
    dataset_path = tmp_path / "benchmark.json"
    dataset_path.write_text(json.dumps([{}, {}]), encoding="utf-8")

    with pytest.raises(DatasetSchemaError, match="Missing required columns"):
        load_dataset(dataset_path, schema=DATASET_SCHEMA)


def test_load_dataset_skips_empty_rows_yielded_by_parser(monkeypatch, tmp_path):
    dataset_path = tmp_path / "benchmark.json"

    def parser(path, schema):
        yield {}
        yield {"id": 1, "prompt": "Prompt", "expected": "Expected"}

    monkeypatch.setitem(dataset_loader.ROW_PARSERS, ".json", parser)

    rows = load_dataset(dataset_path, schema=DATASET_SCHEMA)

    assert rows == [{"id": 1, "prompt": "Prompt", "expected": "Expected"}]


def test_load_dataset_rejects_docx_without_tables(tmp_path):
    dataset_path = tmp_path / "benchmark.docx"
    Document().save(dataset_path)

    with pytest.raises(DatasetSchemaError, match="at least one table"):
        load_dataset(dataset_path, schema=DATASET_SCHEMA)


def test_load_dataset_rejects_docx_table_without_header(monkeypatch, tmp_path):
    dataset_path = tmp_path / "benchmark.docx"

    class EmptyTable:
        rows = []

    class EmptyDocument:
        tables = [EmptyTable()]

    monkeypatch.setattr(dataset_loader, "Document", lambda path: EmptyDocument())

    with pytest.raises(DatasetSchemaError, match="header row"):
        load_dataset(dataset_path, schema=DATASET_SCHEMA)


def test_load_dataset_skips_empty_docx_rows(tmp_path):
    dataset_path = tmp_path / "benchmark.docx"
    _write_docx_table(
        dataset_path,
        headers=["id", "prompt", "expected"],
        rows=[
            ["1", "First prompt", "First expected"],
            ["", "", ""],
            ["2", "Second prompt", "Second expected"],
        ],
    )

    rows = load_dataset(dataset_path, schema=DATASET_SCHEMA)

    assert rows == [
        {"id": 1, "prompt": "First prompt", "expected": "First expected"},
        {"id": 2, "prompt": "Second prompt", "expected": "Second expected"},
    ]


def test_load_dataset_skips_leading_empty_csv_rows_and_blank_headers(tmp_path):
    dataset_path = tmp_path / "benchmark.csv"
    dataset_path.write_text(
        "\n\n"
        "id,prompt,expected,\n"
        "1,Prompt,Expected,ignored\n",
        encoding="utf-8",
    )

    rows = load_dataset(dataset_path, schema=DATASET_SCHEMA)

    assert rows == [{"id": 1, "prompt": "Prompt", "expected": "Expected"}]


def test_load_dataset_rejects_empty_csv_file(tmp_path):
    dataset_path = tmp_path / "benchmark.csv"
    dataset_path.write_text("\n\n", encoding="utf-8")

    with pytest.raises(DatasetSchemaError, match="Missing required columns"):
        load_dataset(dataset_path, schema=DATASET_SCHEMA)


def test_load_dataset_skips_thousands_of_empty_rows_without_crashing(tmp_path):
    dataset_path = tmp_path / "benchmark.csv"
    empty_rows = "\n" * 5_000
    dataset_path.write_text(
        "id,prompt,expected\n"
        "1,First prompt,First expected\n"
        f"{empty_rows}"
        "2,Second prompt,Second expected\n",
        encoding="utf-8",
    )

    rows = load_dataset(dataset_path, schema=DATASET_SCHEMA)

    assert rows == [
        {"id": 1, "prompt": "First prompt", "expected": "First expected"},
        {"id": 2, "prompt": "Second prompt", "expected": "Second expected"},
    ]
