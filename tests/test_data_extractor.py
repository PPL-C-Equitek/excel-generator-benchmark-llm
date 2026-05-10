import zipfile
from pathlib import Path

import pytest
from docx import Document

import src.data_extractor as data_extractor
from src.data_extractor import (
    _safe_xml_from_bytes,
    _xlsx_column_index,
    extract_text_from_file,
)


def test_extract_text_from_file_reads_csv_and_txt(tmp_path):
    csv_path = tmp_path / "sample.csv"
    txt_path = tmp_path / "sample.txt"
    csv_path.write_text("unit,item,value\nIT,Laptop,15000000\n", encoding="utf-8")
    txt_path.write_text("plain extracted text", encoding="utf-8")

    assert extract_text_from_file(csv_path) == (
        "unit,item,value\nIT,Laptop,15000000\n"
    )
    assert extract_text_from_file(txt_path) == "plain extracted text"


def test_extract_text_from_file_extracts_docx_paragraphs_and_tables(tmp_path):
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Quarterly budget")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Laptop"
    table.cell(1, 1).text = "15000000"
    document.save(docx_path)

    extracted_text = extract_text_from_file(docx_path)

    assert extracted_text.splitlines() == [
        "Quarterly budget",
        "Item\tValue",
        "Laptop\t15000000",
    ]


def test_extract_text_from_file_extracts_pdf_with_pdfplumber(tmp_path, mocker):
    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(b"%PDF mocked bytes")
    pdf_reader = mocker.MagicMock()
    pdf_reader.__enter__.return_value.pages = [
        mocker.Mock(extract_text=mocker.Mock(return_value="First page")),
        mocker.Mock(extract_text=mocker.Mock(return_value=None)),
        mocker.Mock(extract_text=mocker.Mock(return_value="Second page")),
    ]
    pdf_open = mocker.patch(
        "src.data_extractor.pdfplumber.open",
        return_value=pdf_reader,
    )

    extracted_text = extract_text_from_file(pdf_path)

    pdf_open.assert_called_once_with(pdf_path)
    assert extracted_text == "First page\nSecond page"


def test_extract_text_from_file_extracts_png_with_pytesseract(tmp_path, mocker):
    png_path = tmp_path / "invoice.png"
    png_path.write_bytes(b"mock image bytes")
    image = mocker.MagicMock()
    processed_image = mocker.MagicMock()
    processed_image.info = {}
    image.convert.return_value = processed_image
    image_open = mocker.patch(
        "src.data_extractor.Image.open",
        return_value=image,
    )
    image_to_string = mocker.patch(
        "src.data_extractor.pytesseract.image_to_string",
        return_value="Invoice OCR text",
    )

    extracted_text = extract_text_from_file(png_path)

    image_open.assert_called_once_with(png_path)
    image.convert.assert_called_once_with("L")
    assert processed_image.info["dpi"] == (300, 300)
    image_to_string.assert_called_once_with(processed_image)
    assert extracted_text == "Invoice OCR text"


def test_extract_text_from_file_raises_clear_error_when_ocr_is_unavailable(
    tmp_path,
    mocker,
):
    png_path = tmp_path / "invoice.png"
    png_path.write_bytes(b"mock image bytes")
    image = mocker.MagicMock()
    processed_image = mocker.MagicMock()
    processed_image.info = {}
    image.convert.return_value = processed_image
    mocker.patch("src.data_extractor.Image.open", return_value=image)
    mocker.patch(
        "src.data_extractor.pytesseract.image_to_string",
        side_effect=RuntimeError("tesseract is not installed"),
    )

    with pytest.raises(ValueError, match="Tesseract OCR is not available"):
        extract_text_from_file(png_path)


def test_extract_text_from_file_raises_path_hint_for_tesseract_not_found(
    tmp_path,
    monkeypatch,
    mocker,
):
    png_path = tmp_path / "invoice.png"
    png_path.write_bytes(b"mock image bytes")
    image = mocker.MagicMock()
    processed_image = mocker.MagicMock()
    processed_image.info = {}
    image.convert.return_value = processed_image
    mocker.patch("src.data_extractor.Image.open", return_value=image)

    class FakeTesseractNotFoundError(Exception):
        pass

    fake_pytesseract = type(
        "FakePytesseract",
        (),
        {
            "TesseractNotFoundError": FakeTesseractNotFoundError,
            "image_to_string": staticmethod(
                lambda _img: (_ for _ in ()).throw(FakeTesseractNotFoundError())
            ),
        },
    )()
    monkeypatch.setattr(data_extractor, "pytesseract", fake_pytesseract)

    with pytest.raises(ValueError, match="ensure `tesseract` is on PATH"):
        extract_text_from_file(png_path)


def test_extract_text_from_file_rejects_unsupported_extensions(tmp_path):
    unsupported_path = tmp_path / "sample.gif"
    unsupported_path.write_text("unsupported", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported file format"):
        extract_text_from_file(unsupported_path)


def test_extract_text_from_file_requires_pdfplumber_for_pdf(
    tmp_path,
    monkeypatch,
):
    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(b"%PDF mocked bytes")
    monkeypatch.setattr(
        data_extractor,
        "pdfplumber",
        data_extractor._MissingPdfPlumber(),
    )

    with pytest.raises(ValueError, match="pdfplumber is required"):
        extract_text_from_file(pdf_path)


def test_extract_text_from_file_requires_pillow_for_png(tmp_path, monkeypatch):
    png_path = tmp_path / "invoice.png"
    png_path.write_bytes(b"mock image bytes")
    monkeypatch.setattr(
        data_extractor,
        "Image",
        data_extractor._MissingImage(),
    )

    with pytest.raises(ValueError, match="Pillow is required"):
        extract_text_from_file(png_path)


def test_extract_text_from_file_requires_pytesseract_for_png(
    tmp_path,
    monkeypatch,
    mocker,
):
    png_path = tmp_path / "invoice.png"
    png_path.write_bytes(b"mock image bytes")
    monkeypatch.setattr(
        data_extractor,
        "Image",
        mocker.Mock(open=mocker.Mock(return_value=mocker.Mock())),
    )
    monkeypatch.setattr(
        data_extractor,
        "pytesseract",
        data_extractor._MissingTesseract(),
    )

    with pytest.raises(ValueError, match="Tesseract OCR is not available"):
        extract_text_from_file(png_path)


def test_preprocess_png_for_ocr_applies_grayscale_and_dpi(mocker):
    image = mocker.MagicMock()
    processed = mocker.MagicMock()
    processed.info = {}
    image.convert.return_value = processed

    result = data_extractor._preprocess_png_for_ocr(image)

    image.convert.assert_called_once_with("L")
    assert result is processed
    assert processed.info["dpi"] == (300, 300)


def test_preprocess_png_for_ocr_handles_images_without_convert():
    class MinimalImage:
        def __init__(self) -> None:
            self.info: dict[str, object] = {}

    image = MinimalImage()

    result = data_extractor._preprocess_png_for_ocr(image)

    assert result is image
    assert result.info["dpi"] == (300, 300)


def test_close_image_if_possible_ignores_objects_without_close():
    data_extractor._close_image_if_possible(object())


def test_extractor_registry_contains_all_supported_extensions():
    registry = data_extractor._extractor_registry()

    assert set(registry) == set(data_extractor.SUPPORTED_FILE_EXTENSIONS)
    assert registry[".csv"] is data_extractor._extract_text_file


def test_extract_text_from_file_extracts_xlsx_without_shared_strings(tmp_path):
    workbook_path = tmp_path / "inline.xlsx"
    _write_minimal_xlsx_without_shared_strings(workbook_path)

    extracted_text = extract_text_from_file(workbook_path)

    assert extracted_text.splitlines() == [
        "SHEET: Sheet1",
        "Inline Item\t",
    ]


def test_extract_text_from_file_extracts_xlsx_shared_strings_and_numbers(
    tmp_path,
):
    workbook_path = tmp_path / "sample.xlsx"
    with zipfile.ZipFile(workbook_path, "w") as workbook:
        workbook.writestr(
            "xl/workbook.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
            <workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
                xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
              <sheets>
                <sheet name="Sheet1" sheetId="1" r:id="rId1"/>
              </sheets>
            </workbook>""",
        )
        workbook.writestr(
            "xl/_rels/workbook.xml.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
            <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
              <Relationship Id="rId1"
                Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet"
                Target="/xl/worksheets/sheet1.xml"/>
            </Relationships>""",
        )
        workbook.writestr(
            "xl/sharedStrings.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
            <sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
              <si><t>Item</t></si>
              <si><t>Total</t></si>
              <si><t>Laptop</t></si>
            </sst>""",
        )
        workbook.writestr(
            "xl/worksheets/sheet1.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
            <worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
              <sheetData>
                <row r="1">
                  <c r="A1" t="s"><v>0</v></c>
                  <c r="B1" t="s"><v>1</v></c>
                </row>
                <row r="2">
                  <c r="A2" t="s"><v>2</v></c>
                  <c r="B2"><v>15000000</v></c>
                </row>
              </sheetData>
            </worksheet>""",
        )

    extracted_text = extract_text_from_file(workbook_path)

    assert extracted_text.splitlines() == [
        "SHEET: Sheet1",
        "Item\tTotal",
        "Laptop\t15000000",
    ]


def _write_minimal_xlsx_without_shared_strings(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as workbook:
        workbook.writestr(
            "xl/workbook.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
            <workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
                xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
              <sheets>
                <sheet name="Sheet1" sheetId="1" r:id="rId1"/>
              </sheets>
            </workbook>""",
        )
        workbook.writestr(
            "xl/_rels/workbook.xml.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
            <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
              <Relationship Id="rId1"
                Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet"
                Target="worksheets/sheet1.xml"/>
            </Relationships>""",
        )
        workbook.writestr(
            "xl/worksheets/sheet1.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
            <worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
              <sheetData>
                <row r="1">
                  <c r="A1" t="inlineStr"><is><t>Inline Item</t></is></c>
                  <c r="B1"></c>
                </row>
                <row r="2"></row>
              </sheetData>
            </worksheet>""",
        )


@pytest.mark.parametrize(
    ("cell_reference", "expected_index"),
    [("A1", 1), ("Z9", 26), ("AA12", 27), ("BC99", 55)],
)
def test_xlsx_column_index(cell_reference, expected_index):
    assert _xlsx_column_index(cell_reference) == expected_index


def test_missing_extractor_fallbacks_raise_clear_errors():
    with pytest.raises(ValueError, match="pdfplumber is required"):
        data_extractor._MissingPdfPlumber.open(Path("/no/such.pdf"))

    with pytest.raises(ValueError, match="Pillow is required"):
        data_extractor._MissingImage.open(Path("/no/such.png"))

    with pytest.raises(RuntimeError, match="pytesseract is required"):
        data_extractor._MissingTesseract.image_to_string(object())


def test_safe_xml_from_bytes_blocks_doctype_and_entity():
    with pytest.raises(ValueError, match="Unsafe XML declaration"):
        _safe_xml_from_bytes(b"<!DOCTYPE foo><root/>")

    with pytest.raises(ValueError, match="Unsafe XML declaration"):
        _safe_xml_from_bytes(b"<!ENTITY xxe SYSTEM 'file:///etc/passwd'><root/>")


def test_safe_xml_from_bytes_allows_regular_xml():
    root = _safe_xml_from_bytes(b"<root><node>ok</node></root>")

    assert root.tag == "root"
    assert root.find("node").text == "ok"


def test_configure_tesseract_binary_uses_valid_custom_env(monkeypatch, tmp_path):
    custom_binary = tmp_path / "tesseract.exe"
    custom_binary.write_text("", encoding="utf-8")
    fake_pytesseract = type(
        "FakePytesseract",
        (),
        {"pytesseract": type("Inner", (), {"tesseract_cmd": ""})()},
    )()
    monkeypatch.setattr(data_extractor, "pytesseract", fake_pytesseract)
    monkeypatch.setattr(
        data_extractor,
        "shutil",
        type("S", (), {"which": staticmethod(lambda _: None)})(),
    )
    monkeypatch.setenv("TESSERACT_CMD", str(custom_binary))

    data_extractor._configure_tesseract_binary()

    assert fake_pytesseract.pytesseract.tesseract_cmd == str(custom_binary.resolve())


def test_configure_tesseract_binary_ignores_invalid_custom_env(monkeypatch):
    fake_pytesseract = type(
        "FakePytesseract",
        (),
        {"pytesseract": type("Inner", (), {"tesseract_cmd": "unchanged"})()},
    )()
    monkeypatch.setattr(data_extractor, "pytesseract", fake_pytesseract)
    monkeypatch.setattr(
        data_extractor,
        "shutil",
        type("S", (), {"which": staticmethod(lambda _: None)})(),
    )
    monkeypatch.setenv("TESSERACT_CMD", "not-found.exe")

    data_extractor._configure_tesseract_binary()

    assert fake_pytesseract.pytesseract.tesseract_cmd == "unchanged"


def test_configure_tesseract_binary_returns_when_pytesseract_missing(monkeypatch):
    monkeypatch.setattr(data_extractor, "pytesseract", data_extractor._MissingTesseract())

    # Should not raise and should exit early.
    data_extractor._configure_tesseract_binary()


def test_configure_tesseract_binary_returns_when_tesseract_in_path(monkeypatch):
    fake_pytesseract = type(
        "FakePytesseract",
        (),
        {"pytesseract": type("Inner", (), {"tesseract_cmd": "unchanged"})()},
    )()
    monkeypatch.setattr(data_extractor, "pytesseract", fake_pytesseract)
    monkeypatch.setattr(
        data_extractor,
        "shutil",
        type("S", (), {"which": staticmethod(lambda _: "C:/bin/tesseract.exe")})(),
    )
    monkeypatch.delenv("TESSERACT_CMD", raising=False)

    data_extractor._configure_tesseract_binary()

    assert fake_pytesseract.pytesseract.tesseract_cmd == "unchanged"


def test_configure_tesseract_binary_uses_windows_candidate_when_available(
    monkeypatch,
    tmp_path,
):
    candidate_binary = tmp_path / "tesseract.exe"
    candidate_binary.write_text("", encoding="utf-8")
    fake_pytesseract = type(
        "FakePytesseract",
        (),
        {"pytesseract": type("Inner", (), {"tesseract_cmd": ""})()},
    )()
    monkeypatch.setattr(data_extractor, "pytesseract", fake_pytesseract)
    monkeypatch.setattr(
        data_extractor,
        "shutil",
        type("S", (), {"which": staticmethod(lambda _: None)})(),
    )
    monkeypatch.setattr(
        data_extractor,
        "WINDOWS_TESSERACT_CANDIDATES",
        (candidate_binary,),
    )
    monkeypatch.delenv("TESSERACT_CMD", raising=False)

    data_extractor._configure_tesseract_binary()

    assert fake_pytesseract.pytesseract.tesseract_cmd == str(candidate_binary)
