"""CLI entrypoint for running batch LLM benchmark examples."""

from __future__ import annotations

import csv
import json
import os
import re
import sys
import time
import zipfile
from collections.abc import Iterable
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from dotenv import load_dotenv
from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:  # pragma: no cover
    sys.path.insert(0, str(PROJECT_ROOT))  # pragma: no cover

from src.runner import BenchmarkRunner  # noqa: E402


DEFAULT_MODEL_NAMES = (
    "claude-sonnet-4-6",
    "gpt-5.2-codex",
    "gemini-3.1-pro-preview",
    "deepseek-v3-2",
)
DEFAULT_EXAMPLE_DIRS = (
    "synthetic.examples",
    "synthetic.examples.lanjutan",
    "like-real.examples",
)
DEFAULT_DATA_DIR = "data"
DEFAULT_REPORT_DIR = "benchmark_reports"
DEFAULT_RUNTIME_DATASET_DIR = f"{DEFAULT_DATA_DIR}/benchmark_runtime_datasets"
OVERALL_REPORT_FILENAME = "overall_benchmark_report.csv"
CSV_ENCODING = "utf-8-sig"
JSON_ENCODING = "utf-8"
SUPPORTED_INPUT_EXTENSIONS = frozenset({".csv", ".txt", ".docx", ".xlsx"})
REPORT_HEADERS = ["unit", "item", "num_type", "status_type", "value"]
OVERALL_REPORT_FIELDNAMES = [
    "section",
    "case_id",
    "model",
    "best_model",
    "best_score",
    "fastest_model",
    "fastest_seconds",
    "average_score",
    "average_seconds",
    "total_runs",
    "completed_runs",
]
XML_NS = {
    "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    "office_rel": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}

SYSTEM_PROMPT = """
You are a document parsing assistant.

Return ONLY a valid JSON object with no markdown, no code fences, and no
explanation.

The object must have exactly three top-level keys: document_info, summary, and
content_data.

document_info must be an object with source_type (must be exactly "Excel" or
"PDF", case-sensitive) and filename (non-empty string).

summary must be an object with non-empty string keys and scalar values only
(string, number, boolean, or null), no nested objects or arrays.

content_data must be a non-empty array of table objects, each with table_name
(non-empty string, unique), headers (non-empty array of unique non-empty
strings), and rows (array of objects where each object keys match headers
exactly with scalar values only, no nested objects or arrays).

For Excel files, if the file contains multiple sheets, each sheet must be
represented as a separate table object in content_data with table_name set to
the sheet name. Do not merge sheets together.

For PDF files, all extracted content must be combined into a single table
object in content_data regardless of page count.

If the input data contains columns that represent categorical groupings such as
department names, regions, or units, unpivot those columns into rows to produce
a normalized long format table where each row represents a single observation,
each categorical column header becomes a value in a new column called "unit",
and its corresponding cell value becomes a separate column called "value", with
all other columns repeated for each unpivoted row.

The resulting long format table must always use these exact column names: unit,
item, num_type, status_type, value. Never use translated or alternative column
names such as "Nilai", "Tipe", "Status", "Item", or any other language
variant.

Exclude any rows where the value is 0 or null after unpivoting.

The following is an example of the REQUIRED FORMAT ONLY. Do NOT return this
example data; always parse and return the actual uploaded file content:
{"document_info":{"source_type":"Excel","filename":"example.xlsx"},"summary":{"total_sheets":1,"total_rows":2,"total_columns":5},"content_data":[{"table_name":"Sheet1","headers":["unit","item","num_type","status_type","value"],"rows":[{"unit":"IT","item":"Laptop","num_type":"cost","status_type":"target","value":15000000},{"unit":"IT","item":"Laptop","num_type":"cost","status_type":"actual","value":14000000}]}]}
""".strip()


@dataclass(frozen=True)
class ExampleCase:
    """Input and expected-output files for one benchmark example."""

    case_id: str
    input_path: Path
    output_path: Path


@dataclass(frozen=True)
class SkippedCase:
    """Example file that cannot be benchmarked by this text CLI."""

    input_path: Path
    reason: str


def main() -> int:
    """Run all supported example files against all configured models.

    Environment variables:
        MODEL_NAMES: Optional comma-separated model list. When unset, all
            default SUMOPOD models are used.
        EXAMPLE_DIRS: Optional comma-separated example folders.
        REPORT_DIR: Optional destination folder for per-run CSV reports.
        BENCHMARK_DATASET_DIR: Optional folder for generated runtime datasets.

    Returns:
        Process exit code. ``0`` means the batch runner completed.
    """
    load_dotenv(dotenv_path=PROJECT_ROOT / ".env")

    model_names = _model_names_from_env()
    example_dirs = _example_dirs_from_env()
    data_dir = _project_path(DEFAULT_DATA_DIR)
    report_dir = _project_path_from_env("REPORT_DIR", DEFAULT_REPORT_DIR)
    runtime_dataset_dir = _project_path_from_env(
        "BENCHMARK_DATASET_DIR",
        DEFAULT_RUNTIME_DATASET_DIR,
    )
    data_dir.mkdir(parents=True, exist_ok=True)
    report_dir.mkdir(parents=True, exist_ok=True)
    runtime_dataset_dir.mkdir(parents=True, exist_ok=True)

    cases, skipped_cases = _discover_example_cases(example_dirs)

    print("========================================")
    print("LLM Benchmark Batch Runner")
    print("========================================")
    print(f"Models     : {', '.join(model_names)}")
    print(f"Examples   : {len(cases)} supported")
    print(f"Skipped    : {len(skipped_cases)} unsupported")
    print(f"Report dir : {report_dir}")
    print("")

    if skipped_cases:
        print("Skipped files")
        print("-------------")
        for skipped_case in skipped_cases:
            relative_path = skipped_case.input_path.relative_to(PROJECT_ROOT)
            print(f"- {relative_path}: {skipped_case.reason}")
        print("")

    batch_results = _run_batch(
        cases=cases,
        model_names=model_names,
        report_dir=report_dir,
        runtime_dataset_dir=runtime_dataset_dir,
    )
    overall_report_path = _write_overall_report(batch_results, report_dir)
    _print_batch_summary(batch_results, overall_report_path)

    return 0


def _model_names_from_env() -> list[str]:
    """Read benchmark model names from ``MODEL_NAMES`` or defaults.

    Returns:
        Ordered list of model names to benchmark.
    """
    raw_model_names = os.getenv("MODEL_NAMES")
    if raw_model_names:
        return _csv_env_values(raw_model_names)

    return list(DEFAULT_MODEL_NAMES)


def _example_dirs_from_env() -> list[Path]:
    """Read example folders from ``EXAMPLE_DIRS`` or defaults.

    Returns:
        Ordered list of absolute example folder paths.
    """
    raw_example_dirs = os.getenv("EXAMPLE_DIRS")
    if raw_example_dirs:
        return [_project_path(value) for value in _csv_env_values(raw_example_dirs)]

    return [_project_path(value) for value in DEFAULT_EXAMPLE_DIRS]


def _csv_env_values(raw_value: str) -> list[str]:
    """Split a comma-separated environment value into non-empty entries.

    Args:
        raw_value: Comma-separated environment value.

    Returns:
        Trimmed non-empty values.
    """
    return [value.strip() for value in raw_value.split(",") if value.strip()]


def _project_path_from_env(env_name: str, default: str) -> Path:
    """Resolve a path environment variable relative to the project root.

    Args:
        env_name: Environment variable name.
        default: Default relative path.

    Returns:
        Absolute path.
    """
    return _project_path(os.getenv(env_name, default))


def _project_path(raw_path: str | Path) -> Path:
    """Resolve a path relative to the project root.

    Args:
        raw_path: Absolute or project-relative path.

    Returns:
        Absolute path.
    """
    path = Path(raw_path)
    if path.is_absolute():
        return path

    return PROJECT_ROOT / path


def _discover_example_cases(
    example_dirs: list[Path],
) -> tuple[list[ExampleCase], list[SkippedCase]]:
    """Discover supported input/output pairs in example folders.

    Args:
        example_dirs: Folders that contain ``*_input.*`` files.

    Returns:
        Supported cases and skipped unsupported inputs.
    """
    cases: list[ExampleCase] = []
    skipped_cases: list[SkippedCase] = []

    for example_dir in example_dirs:
        for input_path in sorted(example_dir.glob("*_input.*")):
            if input_path.suffix.lower() not in SUPPORTED_INPUT_EXTENSIONS:
                skipped_cases.append(
                    SkippedCase(
                        input_path=input_path,
                        reason=f"unsupported input format {input_path.suffix}",
                    )
                )
                continue

            output_path = _matching_output_path(input_path)
            if output_path is None:
                skipped_cases.append(
                    SkippedCase(
                        input_path=input_path,
                        reason="matching *_output.csv file was not found",
                    )
                )
                continue

            cases.append(
                ExampleCase(
                    case_id=_case_id(input_path),
                    input_path=input_path,
                    output_path=output_path,
                )
            )

    return cases, skipped_cases


def _matching_output_path(input_path: Path) -> Path | None:
    """Find the expected output CSV for an input file.

    Args:
        input_path: Example input file.

    Returns:
        Matching output CSV path, or ``None`` when no match exists.
    """
    prefix = input_path.stem.removesuffix("_input")
    extension_name = input_path.suffix.lower().lstrip(".")
    candidates = [
        input_path.with_name(f"{prefix}_{extension_name}_output.csv"),
        input_path.with_name(f"{prefix}_output.csv"),
    ]

    for candidate in candidates:
        if candidate.exists():
            return candidate

    return None


def _case_id(input_path: Path) -> str:
    """Build a stable case id for report filenames.

    Args:
        input_path: Example input file.

    Returns:
        Stable case identifier.
    """
    folder_name = input_path.parent.name.replace(".", "_")
    return f"{folder_name}__{input_path.stem.removesuffix('_input')}"


def _run_batch(
    *,
    cases: list[ExampleCase],
    model_names: list[str],
    report_dir: Path,
    runtime_dataset_dir: Path,
) -> list[dict[str, Any]]:
    """Run every supported example case against every model.

    Args:
        cases: Example cases to benchmark.
        model_names: Model names to run.
        report_dir: Destination folder for result CSVs.
        runtime_dataset_dir: Destination folder for generated datasets.

    Returns:
        Lightweight batch summaries for terminal output.
    """
    batch_results: list[dict[str, Any]] = []
    total_runs = len(cases) * len(model_names)
    run_index = 0

    for model_name in model_names:
        for example_case in cases:
            run_index += 1
            print(
                f"[{run_index}/{total_runs}] "
                f"{model_name} -> {example_case.case_id}"
            )
            result = _run_single_case(
                example_case=example_case,
                model_name=model_name,
                report_dir=report_dir,
                runtime_dataset_dir=runtime_dataset_dir,
            )
            batch_results.append(result)
            print(
                "  "
                f"status={result['status']} "
                f"score={result['average_score']:.4f} "
                f"report={result['report_path']}"
            )

    return batch_results


def _run_single_case(
    *,
    example_case: ExampleCase,
    model_name: str,
    report_dir: Path,
    runtime_dataset_dir: Path,
) -> dict[str, Any]:
    """Run one example case against one model.

    Args:
        example_case: Example case to benchmark.
        model_name: Model name to run.
        report_dir: Destination folder for the CSV report.
        runtime_dataset_dir: Destination folder for the generated dataset.

    Returns:
        Batch summary row.
    """
    safe_model_name = _safe_filename(model_name)
    runtime_dataset_path = (
        runtime_dataset_dir / f"{safe_model_name}__{example_case.case_id}.json"
    )
    report_path = report_dir / f"{safe_model_name}__{example_case.case_id}.csv"
    prompt = _build_prompt(
        example_case.input_path,
        _read_source_text(example_case.input_path),
    )
    expected_output = _expected_output_from_ground_truth(
        example_case.input_path,
        example_case.output_path,
    )
    _write_runtime_dataset(runtime_dataset_path, prompt, expected_output)

    runner = BenchmarkRunner(
        dataset_path=runtime_dataset_path,
        model=model_name,
        report_path=report_path,
    )
    start_time = time.perf_counter()
    summary = runner.run()
    elapsed_seconds = time.perf_counter() - start_time
    return {
        "model": model_name,
        "case_id": example_case.case_id,
        "input_path": str(example_case.input_path.relative_to(PROJECT_ROOT)),
        "output_path": str(example_case.output_path.relative_to(PROJECT_ROOT)),
        "report_path": str(report_path.relative_to(PROJECT_ROOT)),
        "elapsed_seconds": elapsed_seconds,
        **summary,
    }


def _safe_filename(value: str) -> str:
    """Convert a model name or case id into a safe filename segment.

    Args:
        value: Raw value.

    Returns:
        Filesystem-safe value.
    """
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", value).strip("_")


def _read_source_text(path: Path) -> str:
    """Read a supported source file into text for an LLM prompt.

    Args:
        path: Source file path.

    Returns:
        Text extracted from the source file.

    Raises:
        ValueError: If the source format is not supported by this CLI.
    """
    extension = path.suffix.lower()
    if extension in {".csv", ".txt"}:
        return path.read_text(encoding=CSV_ENCODING)

    if extension == ".docx":
        return _read_docx_text(path)

    if extension == ".xlsx":
        return _read_xlsx_text(path)

    raise ValueError(f"Unsupported input format for main.py: {extension}.")


def _read_docx_text(path: Path) -> str:
    """Extract paragraph and table text from a DOCX file.

    Args:
        path: DOCX file path.

    Returns:
        Plain-text representation of the document.
    """
    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append("\t".join(cells))

    return "\n".join(chunks)


def _read_xlsx_text(path: Path) -> str:
    """Extract worksheet text from an XLSX file using the standard library.

    Args:
        path: XLSX file path.

    Returns:
        Plain-text representation of workbook sheets.
    """
    with zipfile.ZipFile(path) as workbook:
        shared_strings = _xlsx_shared_strings(workbook)
        sheets = _xlsx_sheet_paths(workbook)
        chunks: list[str] = []

        for sheet_name, sheet_path in sheets:
            chunks.append(f"SHEET: {sheet_name}")
            chunks.extend(_xlsx_sheet_rows(workbook, sheet_path, shared_strings))

    return "\n".join(chunks)


def _xlsx_shared_strings(workbook: zipfile.ZipFile) -> list[str]:
    """Read XLSX shared strings.

    Args:
        workbook: Open XLSX zip file.

    Returns:
        Shared string table values.
    """
    if "xl/sharedStrings.xml" not in workbook.namelist():
        return []

    root = ElementTree.fromstring(workbook.read("xl/sharedStrings.xml"))
    strings: list[str] = []
    for item in root.findall("main:si", XML_NS):
        text_parts = [
            text_node.text or ""
            for text_node in item.findall(".//main:t", XML_NS)
        ]
        strings.append("".join(text_parts))

    return strings


def _xlsx_sheet_paths(workbook: zipfile.ZipFile) -> list[tuple[str, str]]:
    """Read sheet names and XML paths from an XLSX workbook.

    Args:
        workbook: Open XLSX zip file.

    Returns:
        Ordered pairs of sheet name and XML path.
    """
    workbook_root = ElementTree.fromstring(workbook.read("xl/workbook.xml"))
    rel_root = ElementTree.fromstring(
        workbook.read("xl/_rels/workbook.xml.rels")
    )
    relationships = {
        relationship.attrib["Id"]: relationship.attrib["Target"]
        for relationship in rel_root.findall("rel:Relationship", XML_NS)
    }
    sheets: list[tuple[str, str]] = []

    for sheet in workbook_root.findall("main:sheets/main:sheet", XML_NS):
        relationship_id = sheet.attrib[
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        ]
        target = relationships[relationship_id].lstrip("/")
        sheet_path = target if target.startswith("xl/") else f"xl/{target}"
        sheets.append((sheet.attrib["name"], sheet_path))

    return sheets


def _xlsx_sheet_rows(
    workbook: zipfile.ZipFile,
    sheet_path: str,
    shared_strings: list[str],
) -> list[str]:
    """Extract tab-separated rows from one XLSX sheet.

    Args:
        workbook: Open XLSX zip file.
        sheet_path: XML path for one sheet.
        shared_strings: Shared string table values.

    Returns:
        Tab-separated sheet rows.
    """
    root = ElementTree.fromstring(workbook.read(sheet_path))
    rows: list[str] = []

    for row in root.findall(".//main:row", XML_NS):
        values_by_column: dict[int, str] = {}
        for cell in row.findall("main:c", XML_NS):
            column_index = _xlsx_column_index(cell.attrib.get("r", ""))
            values_by_column[column_index] = _xlsx_cell_value(
                cell,
                shared_strings,
            )

        if not values_by_column:
            continue

        max_column = max(values_by_column)
        values = [
            values_by_column.get(column_index, "")
            for column_index in range(1, max_column + 1)
        ]
        if any(value.strip() for value in values):
            rows.append("\t".join(values))

    return rows


def _xlsx_cell_value(
    cell: ElementTree.Element,
    shared_strings: list[str],
) -> str:
    """Extract a scalar cell value from XLSX XML.

    Args:
        cell: XLSX cell XML element.
        shared_strings: Shared string table values.

    Returns:
        Cell value as text.
    """
    cell_type = cell.attrib.get("t")
    if cell_type == "inlineStr":
        return "".join(
            text_node.text or ""
            for text_node in cell.findall(".//main:t", XML_NS)
        )

    value_node = cell.find("main:v", XML_NS)
    if value_node is None or value_node.text is None:
        return ""

    if cell_type == "s":
        return shared_strings[int(value_node.text)]

    return value_node.text


def _xlsx_column_index(cell_reference: str) -> int:
    """Convert an XLSX cell reference into a one-based column index.

    Args:
        cell_reference: Cell reference such as ``A1`` or ``BC12``.

    Returns:
        One-based column index.
    """
    column_letters = "".join(
        character for character in cell_reference if character.isalpha()
    )
    column_index = 0
    for character in column_letters:
        column_index = (column_index * 26) + ord(character.upper()) - 64

    return column_index


def _read_ground_truth_rows(path: Path) -> list[dict[str, Any]]:
    """Read normalized ground-truth rows from a CSV file.

    Args:
        path: Ground-truth CSV file path.

    Returns:
        List of normalized row dictionaries.
    """
    rows: list[dict[str, Any]] = []
    with path.open("r", newline="", encoding=CSV_ENCODING) as file:
        for row in csv.DictReader(file):
            clean_row = {key.strip(): value.strip() for key, value in row.items()}
            clean_row["value"] = int(clean_row["value"])
            rows.append(clean_row)

    return rows


def _expected_output_from_ground_truth(
    input_path: Path,
    output_path: Path,
) -> dict[str, Any]:
    """Build expected output matching the main document-parser schema.

    Args:
        input_path: Source file used in the prompt.
        output_path: Ground-truth CSV file.

    Returns:
        Expected JSON object for strict evaluation.
    """
    expected_rows = _read_ground_truth_rows(output_path)
    return {
        "document_info": {
            "source_type": _source_type(input_path),
            "filename": input_path.name,
        },
        "summary": {
            "total_sheets": 1,
            "total_rows": len(expected_rows),
            "total_columns": len(REPORT_HEADERS),
        },
        "content_data": [
            {
                "table_name": "Sheet1",
                "headers": REPORT_HEADERS,
                "rows": expected_rows,
            }
        ],
    }


def _source_type(path: Path) -> str:
    """Map a source path to prompt source-type labels.

    Args:
        path: Source file path.

    Returns:
        ``PDF`` for PDF inputs, otherwise ``Excel``.
    """
    if path.suffix.lower() == ".pdf":
        return "PDF"

    return "Excel"


def _build_prompt(input_path: Path, source_text: str) -> str:
    """Build the final LLM prompt for one benchmark input.

    Args:
        input_path: Source file path.
        source_text: Text extracted from the source file.

    Returns:
        Prompt sent to the LLM.
    """
    return f"""
{SYSTEM_PROMPT}

SOURCE FILENAME:
{input_path.name}

SOURCE:
{source_text}
""".strip()


def _write_runtime_dataset(
    path: Path,
    prompt: str,
    expected_output: dict[str, Any],
) -> None:
    """Write a one-row runtime dataset used by ``BenchmarkRunner``.

    Args:
        path: Runtime dataset destination.
        prompt: Prompt generated from the source input.
        expected_output: Expected output dictionary used for evaluation.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "rows": [
            {
                "prompt": prompt,
                "expected_output": expected_output,
            }
        ]
    }
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding=JSON_ENCODING,
    )


def _write_overall_report(
    batch_results: list[dict[str, Any]],
    report_dir: Path,
) -> Path:
    """Write the overall benchmark comparison report.

    The report contains per-file winners, per-model averages, and overall
    winner rows for best average score and fastest average runtime.

    Args:
        batch_results: Batch result dictionaries.
        report_dir: Destination report folder.

    Returns:
        Path to the written aggregate CSV report.
    """
    report_path = report_dir / OVERALL_REPORT_FILENAME
    rows = [
        *_case_comparison_rows(batch_results),
        *_model_summary_rows(batch_results),
        *_overall_winner_rows(batch_results),
    ]

    with report_path.open("w", newline="", encoding=CSV_ENCODING) as file:
        writer = csv.DictWriter(file, fieldnames=OVERALL_REPORT_FIELDNAMES)
        writer.writeheader()
        writer.writerows(rows)

    return report_path


def _case_comparison_rows(
    batch_results: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Build aggregate rows showing best and fastest model per file.

    Args:
        batch_results: Batch result dictionaries.

    Returns:
        Overall report rows with ``section`` set to ``case_comparison``.
    """
    rows: list[dict[str, Any]] = []
    for case_id, case_results in _group_results(batch_results, "case_id").items():
        best_result = _best_score_result(case_results)
        fastest_result = _fastest_result(case_results)
        rows.append(
            {
                "section": "case_comparison",
                "case_id": case_id,
                "model": "",
                "best_model": best_result["model"],
                "best_score": best_result["average_score"],
                "fastest_model": fastest_result["model"],
                "fastest_seconds": fastest_result["elapsed_seconds"],
                "average_score": "",
                "average_seconds": "",
                "total_runs": len(case_results),
                "completed_runs": _completed_run_count(case_results),
            }
        )

    return rows


def _model_summary_rows(
    batch_results: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Build aggregate rows showing average score and runtime per model.

    Args:
        batch_results: Batch result dictionaries.

    Returns:
        Overall report rows with ``section`` set to ``model_summary``.
    """
    rows: list[dict[str, Any]] = []
    for model_name, model_results in _group_results(batch_results, "model").items():
        rows.append(
            {
                "section": "model_summary",
                "case_id": "",
                "model": model_name,
                "best_model": "",
                "best_score": "",
                "fastest_model": "",
                "fastest_seconds": "",
                "average_score": _average_float(
                    result["average_score"] for result in model_results
                ),
                "average_seconds": _average_float(
                    result["elapsed_seconds"] for result in model_results
                ),
                "total_runs": len(model_results),
                "completed_runs": _completed_run_count(model_results),
            }
        )

    return rows


def _overall_winner_rows(
    batch_results: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Build rows for the overall best-score and fastest-runtime models.

    Args:
        batch_results: Batch result dictionaries.

    Returns:
        Overall winner rows.
    """
    model_rows = _model_summary_rows(batch_results)
    if not model_rows:
        return []

    best_average_row = max(
        model_rows,
        key=lambda row: (
            float(row["average_score"]),
            -float(row["average_seconds"]),
        ),
    )
    fastest_average_row = min(
        model_rows,
        key=lambda row: (
            float(row["average_seconds"]),
            -float(row["average_score"]),
        ),
    )
    return [
        {
            "section": "overall_best_average",
            "case_id": "",
            "model": best_average_row["model"],
            "best_model": "",
            "best_score": "",
            "fastest_model": "",
            "fastest_seconds": "",
            "average_score": best_average_row["average_score"],
            "average_seconds": best_average_row["average_seconds"],
            "total_runs": best_average_row["total_runs"],
            "completed_runs": best_average_row["completed_runs"],
        },
        {
            "section": "overall_fastest",
            "case_id": "",
            "model": fastest_average_row["model"],
            "best_model": "",
            "best_score": "",
            "fastest_model": "",
            "fastest_seconds": "",
            "average_score": fastest_average_row["average_score"],
            "average_seconds": fastest_average_row["average_seconds"],
            "total_runs": fastest_average_row["total_runs"],
            "completed_runs": fastest_average_row["completed_runs"],
        },
    ]


def _group_results(
    batch_results: list[dict[str, Any]],
    key: str,
) -> dict[str, list[dict[str, Any]]]:
    """Group batch result dictionaries by a string key.

    Args:
        batch_results: Batch result dictionaries.
        key: Result dictionary key to group by.

    Returns:
        Grouped batch results.
    """
    grouped_results: dict[str, list[dict[str, Any]]] = {}
    for result in batch_results:
        grouped_results.setdefault(str(result[key]), []).append(result)

    return grouped_results


def _best_score_result(results: list[dict[str, Any]]) -> dict[str, Any]:
    """Find the best-scoring result, using runtime as tie-breaker.

    Args:
        results: Batch results for one case.

    Returns:
        Best result dictionary.
    """
    return max(
        results,
        key=lambda result: (
            float(result["average_score"]),
            -float(result["elapsed_seconds"]),
        ),
    )


def _fastest_result(results: list[dict[str, Any]]) -> dict[str, Any]:
    """Find the fastest result, using score as tie-breaker.

    Args:
        results: Batch results for one case.

    Returns:
        Fastest result dictionary.
    """
    return min(
        results,
        key=lambda result: (
            float(result["elapsed_seconds"]),
            -float(result["average_score"]),
        ),
    )


def _completed_run_count(results: list[dict[str, Any]]) -> int:
    """Count completed runs in a result collection.

    Args:
        results: Batch result dictionaries.

    Returns:
        Number of completed runs.
    """
    return sum(1 for result in results if result["status"] == "completed")


def _average_float(values: Iterable[Any]) -> float:
    """Calculate a safe floating-point average.

    Args:
        values: Iterable numeric values.

    Returns:
        Average value, or ``0.0`` when empty.
    """
    numeric_values = [float(value) for value in values]
    if not numeric_values:
        return 0.0

    return sum(numeric_values) / len(numeric_values)


def _print_batch_summary(
    batch_results: list[dict[str, Any]],
    overall_report_path: Path,
) -> None:
    """Print a readable summary table for batch results.

    Args:
        batch_results: Batch result dictionaries.
        overall_report_path: Aggregate report path.
    """
    print("")
    print("Batch Summary")
    print("-------------")
    if not batch_results:
        print("No supported benchmark cases were found.")
        return

    for result in batch_results:
        print(
            f"{result['model']} | {result['case_id']} | "
            f"status={result['status']} | "
            f"score={result['average_score']:.4f} | "
            f"success={result['successful_evaluations']}/"
            f"{result['total_rows']}"
        )

    winner_rows = _overall_winner_rows(batch_results)
    if winner_rows:
        best_average_row, fastest_row = winner_rows
        print("")
        print("Overall Winners")
        print("---------------")
        print(
            "Best average score: "
            f"{best_average_row['model']} "
            f"({float(best_average_row['average_score']):.4f})"
        )
        print(
            "Fastest average runtime: "
            f"{fastest_row['model']} "
            f"({float(fastest_row['average_seconds']):.4f}s)"
        )

    print("")
    print(f"Overall report: {overall_report_path}")


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())
