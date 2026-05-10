"""Local text extraction utilities for benchmark input files."""

from __future__ import annotations

import zipfile
import os
import shutil
from pathlib import Path
from collections.abc import Callable
from typing import Any
from xml.etree import ElementTree

from docx import Document


class _MissingPdfPlumber:
    """Fallback object used when pdfplumber is not installed."""

    @staticmethod
    def open(file_path: Path) -> Any:
        """Raise a clear error for unavailable PDF extraction support."""
        raise ValueError(
            "pdfplumber is required to extract PDF text. "
            "Install it with `pip install pdfplumber`."
        )


class _MissingImage:
    """Fallback object used when Pillow is not installed."""

    @staticmethod
    def open(file_path: Path) -> Any:
        """Raise a clear error for unavailable image loading support."""
        raise ValueError(
            "Pillow is required to open image files. "
            "Install it with `pip install Pillow`."
        )


class _MissingTesseract:
    """Fallback object used when pytesseract is not installed."""

    @staticmethod
    def image_to_string(image: Any) -> str:
        """Raise a clear error for unavailable OCR support."""
        raise RuntimeError(
            "pytesseract is required to extract OCR text from images."
        )


try:
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = _MissingPdfPlumber()

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = _MissingImage()

try:
    import pytesseract
except ImportError:  # pragma: no cover
    pytesseract = _MissingTesseract()


TEXT_ENCODING = "utf-8-sig"
SUPPORTED_FILE_EXTENSIONS = frozenset(
    {".csv", ".txt", ".docx", ".pdf", ".png", ".xlsx"}
)
XML_NS = {
    "main": "http" + "://schemas.openxmlformats.org/spreadsheetml/2006/main",
    "rel": "http" + "://schemas.openxmlformats.org/package/2006/relationships",
}
WINDOWS_TESSERACT_CANDIDATES = (
    Path("C:/Program Files/Tesseract-OCR/tesseract.exe"),
    Path("C:/Program Files (x86)/Tesseract-OCR/tesseract.exe"),
)
DISALLOWED_XML_TOKENS = (b"<!DOCTYPE", b"<!ENTITY")


def _configure_tesseract_binary() -> None:
    """Configure pytesseract binary path for environments without PATH setup."""
    if isinstance(pytesseract, _MissingTesseract):
        return

    custom_binary = os.getenv("TESSERACT_CMD")
    if custom_binary:
        custom_binary_path = Path(custom_binary.strip()).expanduser().resolve()
        if (
            custom_binary_path.exists()
            and custom_binary_path.is_file()
            and custom_binary_path.name.lower().startswith("tesseract")
        ):
            pytesseract.pytesseract.tesseract_cmd = str(custom_binary_path)
        return

    if shutil.which("tesseract"):
        return

    for candidate in WINDOWS_TESSERACT_CANDIDATES:
        if candidate.exists():
            pytesseract.pytesseract.tesseract_cmd = str(candidate)
            return


_configure_tesseract_binary()


def extract_text_from_file(file_path: str | Path) -> str:
    """Extract plain text from a supported benchmark input file.

    Args:
        file_path: Path to a ``.csv``, ``.txt``, ``.docx``, ``.pdf``,
            ``.png``, or ``.xlsx`` file.

    Returns:
        Extracted plain text suitable for inclusion in an LLM prompt.

    Raises:
        ValueError: If the extension is unsupported, an optional extractor
            dependency is missing, or OCR extraction fails.
    """
    path = Path(file_path)
    extension = path.suffix.lower()
    extractor = _extractor_registry().get(extension)
    if extractor is None:
        raise ValueError(f"Unsupported file format for extraction: {extension}.")
    return extractor(path)


def _extract_text_file(path: Path) -> str:
    """Extract plain text from CSV/TXT files."""
    return path.read_text(encoding=TEXT_ENCODING)


def _extract_docx_text(path: Path) -> str:
    """Extract paragraph and table text from a DOCX file.

    Args:
        path: DOCX file path.

    Returns:
        Plain-text document content.
    """
    document = Document(path)
    chunks = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append("\t".join(cells))

    return "\n".join(chunks)


def _extract_pdf_text(path: Path) -> str:
    """Extract text from a PDF file using pdfplumber.

    Args:
        path: PDF file path.

    Returns:
        Newline-joined text extracted from all pages.
    """
    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                chunks.append(page_text)

    return "\n".join(chunks)


def _extract_png_text(path: Path) -> str:
    """Extract OCR text from a PNG image using Pillow and pytesseract.

    Args:
        path: PNG image path.

    Returns:
        OCR text extracted from the image.

    Raises:
        ValueError: If OCR extraction cannot run.
    """
    image = Image.open(path)
    prepared_image = _preprocess_png_for_ocr(image)
    try:
        return pytesseract.image_to_string(prepared_image).strip()
    except Exception as exc:
        tesseract_not_found_error = getattr(
            pytesseract,
            "TesseractNotFoundError",
            None,
        )
        if (
            tesseract_not_found_error is not None
            and isinstance(exc, tesseract_not_found_error)
        ):
            raise ValueError(
                "Tesseract OCR is not available. Install the Tesseract binary "
                "and ensure `tesseract` is on PATH (or set `TESSERACT_CMD`)."
            ) from exc
        raise ValueError(
            "Tesseract OCR is not available. Install the Tesseract binary "
            "and the `pytesseract` Python package before extracting PNG text."
        ) from exc
    finally:
        _close_image_if_possible(prepared_image)
        if prepared_image is not image:
            _close_image_if_possible(image)


def _preprocess_png_for_ocr(image: Any) -> Any:
    """Apply lightweight OCR preprocessing (grayscale + DPI hint)."""
    convert = getattr(image, "convert", None)
    processed_image = convert("L") if callable(convert) else image

    info = getattr(processed_image, "info", None)
    if isinstance(info, dict):
        info.setdefault("dpi", (300, 300))

    return processed_image


def _close_image_if_possible(image: Any) -> None:
    """Close image resources when object exposes a ``close`` method."""
    close_image = getattr(image, "close", None)
    if callable(close_image):
        close_image()


def _extractor_registry() -> dict[str, Callable[[Path], str]]:
    """Return extractor handlers keyed by normalized file extension."""
    return {
        ".csv": _extract_text_file,
        ".txt": _extract_text_file,
        ".docx": _extract_docx_text,
        ".pdf": _extract_pdf_text,
        ".png": _extract_png_text,
        ".xlsx": _extract_xlsx_text,
    }


def _extract_xlsx_text(path: Path) -> str:
    """Extract worksheet text from an XLSX file using the standard library.

    Args:
        path: XLSX file path.

    Returns:
        Plain-text representation of workbook sheets.
    """
    with zipfile.ZipFile(path) as workbook:
        shared_strings = _xlsx_shared_strings(workbook)
        sheets = _xlsx_sheet_paths(workbook)
        chunks: list[str] = []

        for sheet_name, sheet_path in sheets:
            chunks.append(f"SHEET: {sheet_name}")
            chunks.extend(_xlsx_sheet_rows(workbook, sheet_path, shared_strings))

    return "\n".join(chunks)


def _xlsx_shared_strings(workbook: zipfile.ZipFile) -> list[str]:
    """Read XLSX shared strings.

    Args:
        workbook: Open XLSX zip file.

    Returns:
        Shared string table values.
    """
    if "xl/sharedStrings.xml" not in workbook.namelist():
        return []

    root = _safe_xml_from_bytes(workbook.read("xl/sharedStrings.xml"))
    strings: list[str] = []
    for item in root.findall("main:si", XML_NS):
        text_parts = [
            text_node.text or ""
            for text_node in item.findall(".//main:t", XML_NS)
        ]
        strings.append("".join(text_parts))

    return strings


def _xlsx_sheet_paths(workbook: zipfile.ZipFile) -> list[tuple[str, str]]:
    """Read sheet names and XML paths from an XLSX workbook.

    Args:
        workbook: Open XLSX zip file.

    Returns:
        Ordered pairs of sheet name and XML path.
    """
    workbook_root = _safe_xml_from_bytes(workbook.read("xl/workbook.xml"))
    rel_root = _safe_xml_from_bytes(workbook.read("xl/_rels/workbook.xml.rels"))
    relationships = {
        relationship.attrib["Id"]: relationship.attrib["Target"]
        for relationship in rel_root.findall("rel:Relationship", XML_NS)
    }
    sheets: list[tuple[str, str]] = []

    for sheet in workbook_root.findall("main:sheets/main:sheet", XML_NS):
        relationship_id = sheet.attrib[
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        ]
        target = relationships[relationship_id].lstrip("/")
        sheet_path = target if target.startswith("xl/") else f"xl/{target}"
        sheets.append((sheet.attrib["name"], sheet_path))

    return sheets


def _xlsx_sheet_rows(
    workbook: zipfile.ZipFile,
    sheet_path: str,
    shared_strings: list[str],
) -> list[str]:
    """Extract tab-separated rows from one XLSX sheet.

    Args:
        workbook: Open XLSX zip file.
        sheet_path: XML path for one sheet.
        shared_strings: Shared string table values.

    Returns:
        Tab-separated sheet rows.
    """
    root = _safe_xml_from_bytes(workbook.read(sheet_path))
    rows: list[str] = []

    for row in root.findall(".//main:row", XML_NS):
        values_by_column: dict[int, str] = {}
        for cell in row.findall("main:c", XML_NS):
            column_index = _xlsx_column_index(cell.attrib.get("r", ""))
            values_by_column[column_index] = _xlsx_cell_value(
                cell,
                shared_strings,
            )

        if not values_by_column:
            continue

        max_column = max(values_by_column)
        values = [
            values_by_column.get(column_index, "")
            for column_index in range(1, max_column + 1)
        ]
        if any(value.strip() for value in values):
            rows.append("\t".join(values))

    return rows


def _xlsx_cell_value(
    cell: ElementTree.Element,
    shared_strings: list[str],
) -> str:
    """Extract a scalar cell value from XLSX XML.

    Args:
        cell: XLSX cell XML element.
        shared_strings: Shared string table values.

    Returns:
        Cell value as text.
    """
    cell_type = cell.attrib.get("t")
    if cell_type == "inlineStr":
        return "".join(
            text_node.text or ""
            for text_node in cell.findall(".//main:t", XML_NS)
        )

    value_node = cell.find("main:v", XML_NS)
    if value_node is None or value_node.text is None:
        return ""

    if cell_type == "s":
        return shared_strings[int(value_node.text)]

    return value_node.text


def _xlsx_column_index(cell_reference: str) -> int:
    """Convert an XLSX cell reference into a one-based column index.

    Args:
        cell_reference: Cell reference such as ``A1`` or ``BC12``.

    Returns:
        One-based column index.
    """
    column_letters = "".join(
        character for character in cell_reference if character.isalpha()
    )
    column_index = 0
    for character in column_letters:
        column_index = (column_index * 26) + ord(character.upper()) - 64

    return column_index


def _safe_xml_from_bytes(raw_xml: bytes) -> ElementTree.Element:
    """Parse XML bytes after blocking unsafe DTD/entity declarations.

    Args:
        raw_xml: XML bytes loaded from a workbook archive.

    Returns:
        Parsed XML element root.

    Raises:
        ValueError: If XML contains disallowed DTD/entity declarations.
    """
    upper_raw_xml = raw_xml.upper()
    if any(token in upper_raw_xml for token in DISALLOWED_XML_TOKENS):
        raise ValueError("Unsafe XML declaration found in XLSX content.")
    return ElementTree.fromstring(raw_xml)
