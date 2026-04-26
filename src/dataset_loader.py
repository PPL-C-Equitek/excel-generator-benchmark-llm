"""Dataset loading and schema validation utilities.

This module supports benchmark datasets stored as CSV, JSON, or DOCX files.
Rows are parsed through small iterators so blank rows can be skipped before
schema validation and casting.
"""

from __future__ import annotations

import csv
import json
from collections.abc import Callable, Iterable, Iterator, Mapping
from pathlib import Path
from typing import Any

from docx import Document


CSV_EXTENSION = ".csv"
JSON_EXTENSION = ".json"
DOCX_EXTENSION = ".docx"
SUPPORTED_EXTENSIONS = frozenset(
    {CSV_EXTENSION, JSON_EXTENSION, DOCX_EXTENSION}
)

JSON_ROWS_KEY = "rows"
JSON_DATA_KEY = "data"
UTF_8_ENCODING = "utf-8"
UTF_8_SIG_ENCODING = "utf-8-sig"

Row = dict[str, Any]
Schema = Mapping[str, type[Any]]
RowIterator = Iterator[Row]
RowParser = Callable[[Path, Schema], RowIterator]


class DatasetSchemaError(ValueError):
    """Raised when dataset columns or values do not match the schema."""


class UnsupportedDatasetFormatError(ValueError):
    """Raised when a dataset file extension is not supported."""


def load_dataset(path: str | Path, schema: Schema) -> list[Row]:
    """Load a dataset file and validate every row against a schema.

    Args:
        path: File path to a CSV, JSON, or DOCX dataset.
        schema: Mapping of required column names to expected Python types.

    Returns:
        A list of validated and type-cast dataset rows.

    Raises:
        DatasetSchemaError: If the schema is invalid, required columns are
            missing, or row values cannot be cast to the required types.
        UnsupportedDatasetFormatError: If the file extension is not supported.
    """
    dataset_path = Path(path)
    _validate_schema(schema)

    extension = dataset_path.suffix.lower()
    parser = ROW_PARSERS.get(extension)
    if parser is None:
        _raise_unsupported_format(extension, dataset_path)

    rows = parser(dataset_path, schema)
    return list(_validate_and_cast_rows(rows, schema))


def _raise_unsupported_format(extension: str, path: Path) -> None:
    """Raise an unsupported-format exception with the accepted extensions.

    Args:
        extension: Lowercase extension extracted from the dataset path.
        path: Dataset path used in the error message.

    Raises:
        UnsupportedDatasetFormatError: Always raised by this helper.
    """
    supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
    raise UnsupportedDatasetFormatError(
        f"Unsupported dataset format '{extension}' for {path}. "
        f"Supported formats: {supported}."
    )


def _validate_schema(schema: Schema) -> None:
    """Validate that a schema defines required columns and type objects.

    Args:
        schema: Mapping of column names to Python type objects.

    Raises:
        DatasetSchemaError: If the schema is empty or maps a column to a
            non-type value.
    """
    if not schema:
        raise DatasetSchemaError("Schema must define at least one required column.")

    invalid_columns = [
        column
        for column, expected_type in schema.items()
        if not isinstance(expected_type, type)
    ]
    if not invalid_columns:
        return

    joined_columns = ", ".join(invalid_columns)
    raise DatasetSchemaError(
        f"Schema columns must map to Python types: {joined_columns}."
    )


def _iter_csv_rows(path: Path, schema: Schema) -> RowIterator:
    """Yield non-empty CSV rows as dictionaries.

    Args:
        path: CSV dataset path.
        schema: Required column schema used to validate the CSV header.

    Yields:
        Raw row dictionaries keyed by CSV header names.

    Raises:
        DatasetSchemaError: If the header is missing required schema columns.
    """
    with path.open("r", newline="", encoding=UTF_8_SIG_ENCODING) as file:
        reader = csv.reader(file)
        headers = _next_non_empty_csv_row(reader)
        _validate_required_columns(headers, schema)

        for values in reader:
            if _is_empty_sequence(values):
                continue
            yield _row_from_values(headers, values)


def _iter_json_rows(path: Path, schema: Schema) -> RowIterator:
    """Yield non-empty JSON rows as dictionaries.

    The JSON dataset may be a single object, a list of objects, or an object
    containing a ``rows`` or ``data`` list.

    Args:
        path: JSON dataset path.
        schema: Required column schema used to validate the first row.

    Yields:
        Raw row dictionaries from the JSON payload.

    Raises:
        DatasetSchemaError: If the payload is not an object/list of objects,
            or required schema columns are missing.
    """
    with path.open("r", encoding=UTF_8_ENCODING) as file:
        payload = json.load(file)

    records = _json_records_from_payload(payload)
    if not records:
        return

    first_non_empty_index = _first_non_empty_record_index(records)
    headers = _headers_from_records(records, first_non_empty_index)
    _validate_required_columns(headers, schema)

    start_index = first_non_empty_index or 0
    for record in records[start_index:]:
        if not isinstance(record, dict):
            raise DatasetSchemaError("JSON dataset rows must be objects.")
        if _is_empty_mapping(record):
            continue
        yield _normalize_row_keys(record)


def _json_records_from_payload(payload: Any) -> list[Any]:
    """Extract candidate row records from a loaded JSON payload.

    Args:
        payload: JSON value returned by ``json.load``.

    Returns:
        A list of candidate JSON row values.

    Raises:
        DatasetSchemaError: If the JSON payload cannot represent dataset rows.
    """
    if isinstance(payload, dict):
        if JSON_ROWS_KEY in payload:
            return _ensure_json_record_list(payload[JSON_ROWS_KEY])
        if JSON_DATA_KEY in payload:
            return _ensure_json_record_list(payload[JSON_DATA_KEY])
        return [payload]

    return _ensure_json_record_list(payload)


def _ensure_json_record_list(records: Any) -> list[Any]:
    """Ensure a JSON records value is list-like for downstream validation.

    Args:
        records: Candidate value that should contain JSON row objects.

    Returns:
        The records value when it is a list.

    Raises:
        DatasetSchemaError: If records is not a list.
    """
    if isinstance(records, list):
        return records

    raise DatasetSchemaError(
        "JSON dataset must contain an object or list of objects."
    )


def _first_non_empty_record_index(records: list[Any]) -> int | None:
    """Find the first JSON row that contains at least one non-empty value.

    Args:
        records: Candidate JSON row records.

    Returns:
        The zero-based index of the first non-empty row, or ``None`` when all
        rows are empty.

    Raises:
        DatasetSchemaError: If any inspected record is not an object.
    """
    for index, record in enumerate(records):
        if not isinstance(record, dict):
            raise DatasetSchemaError("JSON dataset rows must be objects.")
        if _is_empty_mapping(record):
            continue
        return index

    return None


def _headers_from_records(
    records: list[Any],
    first_non_empty_index: int | None,
) -> list[str]:
    """Read headers from the first non-empty JSON record.

    Args:
        records: Candidate JSON row records.
        first_non_empty_index: Index returned by
            ``_first_non_empty_record_index``.

    Returns:
        Header names from the first non-empty record, or an empty list.
    """
    if first_non_empty_index is None:
        return []

    record = records[first_non_empty_index]
    return [str(header) for header in record.keys()]


def _normalize_row_keys(row: Mapping[str, Any]) -> Row:
    """Normalize row dictionary keys for schema-consistent lookup.

    Args:
        row: Raw row mapping from a supported dataset source.

    Returns:
        A row dictionary with stripped string keys.
    """
    return {_clean_header(key): value for key, value in row.items()}


def _iter_docx_rows(path: Path, schema: Schema) -> RowIterator:
    """Yield non-empty rows from the first table in a DOCX document.

    Args:
        path: DOCX dataset path.
        schema: Required column schema used to validate the table header.

    Yields:
        Raw row dictionaries keyed by table header names.

    Raises:
        DatasetSchemaError: If the document has no table, the first table has
            no header row, or required schema columns are missing.
    """
    document = Document(path)
    if not document.tables:
        raise DatasetSchemaError("DOCX dataset must contain at least one table.")

    table = document.tables[0]
    if not table.rows:
        raise DatasetSchemaError("DOCX dataset table must contain a header row.")

    headers = [_cell_text(cell) for cell in table.rows[0].cells]
    _validate_required_columns(headers, schema)

    for table_row in table.rows[1:]:
        values = [_cell_text(cell) for cell in table_row.cells]
        if _is_empty_sequence(values):
            continue
        yield _row_from_values(headers, values)


def _validate_and_cast_rows(rows: Iterable[Row], schema: Schema) -> RowIterator:
    """Validate required values and cast rows according to the schema.

    Args:
        rows: Parsed raw dataset rows.
        schema: Mapping of required columns to expected Python types.

    Yields:
        Validated rows containing only schema-defined columns.

    Raises:
        DatasetSchemaError: If a required value is empty or cannot be cast to
            the expected type.
    """
    for row_number, row in enumerate(rows, start=1):
        if _is_empty_mapping(row):
            continue

        cast_row: Row = {}
        for column, expected_type in schema.items():
            value = row.get(column)
            if _is_empty_value(value):
                raise DatasetSchemaError(
                    f"Column '{column}' is required and cannot be empty "
                    f"at row {row_number}."
                )
            cast_row[column] = _cast_value(
                column,
                value,
                expected_type,
                row_number,
            )

        yield cast_row


def _cast_value(
    column: str,
    value: Any,
    expected_type: type[Any],
    row_number: int,
) -> Any:
    """Cast a dataset value to the schema's expected type.

    Args:
        column: Column name used in error messages.
        value: Raw value read from the dataset.
        expected_type: Python type required by the schema.
        row_number: One-based row number used in error messages.

    Returns:
        The original value when already typed correctly, or the cast value.

    Raises:
        DatasetSchemaError: If the value cannot be cast to ``expected_type``.
    """
    if isinstance(value, expected_type):
        return value

    if expected_type is str:
        return str(value)

    try:
        return expected_type(value)
    except (TypeError, ValueError) as exc:
        raise DatasetSchemaError(
            f"Column '{column}' must be {expected_type.__name__} "
            f"at row {row_number}."
        ) from exc


def _next_non_empty_csv_row(reader: Iterator[list[str]]) -> list[str]:
    """Return the first non-empty CSV row from a reader.

    Args:
        reader: CSV reader yielding raw row values.

    Returns:
        Cleaned header values from the first non-empty row, or an empty list.
    """
    for row in reader:
        if _is_empty_sequence(row):
            continue
        return [_clean_header(value) for value in row]

    return []


def _validate_required_columns(headers: Iterable[str], schema: Schema) -> None:
    """Validate that parsed headers contain every schema column.

    Args:
        headers: Parsed column names from a dataset source.
        schema: Required column schema.

    Raises:
        DatasetSchemaError: If any schema column is missing.
    """
    normalized_headers = {_clean_header(header) for header in headers}
    missing_columns = [
        column
        for column in schema
        if _clean_header(column) not in normalized_headers
    ]
    if not missing_columns:
        return

    joined_columns = ", ".join(missing_columns)
    raise DatasetSchemaError(f"Missing required columns: {joined_columns}.")


def _row_from_values(headers: list[str], values: list[Any]) -> Row:
    """Build a row dictionary from headers and raw positional values.

    Args:
        headers: Header names for the row.
        values: Raw row values aligned by position with ``headers``.

    Returns:
        A dictionary mapping non-empty headers to row values. Missing trailing
        values are represented as empty strings for validation.
    """
    row: Row = {}
    for index, header in enumerate(headers):
        if not header:
            continue
        row[header] = values[index] if index < len(values) else ""

    return row


def _cell_text(cell: Any) -> str:
    """Return stripped text from a DOCX table cell.

    Args:
        cell: Python-docx table cell object.

    Returns:
        The stripped cell text.
    """
    return cell.text.strip()


def _clean_header(value: Any) -> str:
    """Normalize a column header value.

    Args:
        value: Raw header value from a supported dataset source.

    Returns:
        The header converted to a stripped string.
    """
    return str(value).strip()


def _is_empty_value(value: Any) -> bool:
    """Check whether a scalar dataset value should be treated as empty.

    Args:
        value: Candidate dataset value.

    Returns:
        ``True`` when the value is ``None`` or blank text.
    """
    return value is None or (isinstance(value, str) and value.strip() == "")


def _is_empty_sequence(values: Iterable[Any]) -> bool:
    """Check whether all values in a row-like sequence are empty.

    Args:
        values: Iterable of candidate row values.

    Returns:
        ``True`` when every value is empty.
    """
    return all(_is_empty_value(value) for value in values)


def _is_empty_mapping(row: Mapping[str, Any]) -> bool:
    """Check whether all values in a row mapping are empty.

    Args:
        row: Row mapping keyed by column names.

    Returns:
        ``True`` when every mapped value is empty.
    """
    return all(_is_empty_value(value) for value in row.values())


ROW_PARSERS: Mapping[str, RowParser] = {
    CSV_EXTENSION: _iter_csv_rows,
    JSON_EXTENSION: _iter_json_rows,
    DOCX_EXTENSION: _iter_docx_rows,
}
